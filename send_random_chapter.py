import os
import re
import ssl
import random
import smtplib
from pathlib import Path
from email.message import EmailMessage
from zoneinfo import ZoneInfo
from datetime import datetime

from docx import Document


# Regex per riconoscere una cella che contiene SOLO un numero romano
ROMAN_RE = re.compile(
    r"^(?=[MDCLXVI])M{0,4}(CM|CD|D?C{0,3})(XC|XL|L?X{0,3})(IX|IV|V?I{0,3})$"
)


# Nome del file DOCX.
# Ti consiglio di rinominare il file caricato in: spigolature.docx
DOCX_PATH = Path(os.getenv("DOCX_PATH", "spigolature.docx"))

# File in cui vengono salvati i capitoli già inviati
HISTORY_PATH = Path(os.getenv("HISTORY_PATH", "cronologia.txt"))

# Numero atteso di capitoli nel documento
EXPECTED_CHAPTER_COUNT = int(os.getenv("EXPECTED_CHAPTER_COUNT", "166"))

# Configurazione Gmail SMTP
SMTP_SERVER = "smtp.gmail.com"
SMTP_PORT = 465


def should_run_now() -> bool:
    """
    Su GitHub Actions:
    - se il workflow viene lanciato manualmente, invia sempre;
    - se parte da cron, invia solo alle ore locali previste.

    Questo serve perché GitHub Actions usa UTC,
    mentre noi vogliamo inviare alle 6:00 e alle 19:00 ora italiana.
    """
    if os.getenv("GITHUB_EVENT_NAME") == "workflow_dispatch":
        return True

    enforce = os.getenv("ENFORCE_LOCAL_HOURS", "false").lower() == "true"

    if not enforce:
        return True

    timezone = os.getenv("LOCAL_TZ", "Europe/Rome")
    allowed_hours_raw = os.getenv("ALLOWED_LOCAL_HOURS", "6,19")

    allowed_hours = {
        int(hour.strip())
        for hour in allowed_hours_raw.split(",")
        if hour.strip()
    }

    now = datetime.now(ZoneInfo(timezone))

    return now.hour in allowed_hours


def clean_cell_text(cell) -> str:
    """
    Estrae il testo da una cella del DOCX.

    Gestisce:
    - celle vuote;
    - righe vuote;
    - celle con soli spazi;
    - più paragrafi dentro la stessa cella.
    """
    paragraphs = []

    for paragraph in cell.paragraphs:
        text = paragraph.text.strip()

        if text:
            paragraphs.append(text)

    return "\n".join(paragraphs).strip()


def load_chapters(docx_path: Path) -> list[dict]:
    """
    Carica il DOCX, legge doc.tables[0],
    trova i capitoli tramite numero romano
    ed estrae il testo di ogni capitolo.
    """
    if not docx_path.exists():
        raise FileNotFoundError(f"File DOCX non trovato: {docx_path}")

    doc = Document(str(docx_path))

    if not doc.tables:
        raise ValueError("Il documento non contiene tabelle.")

    table = doc.tables[0]

    chapter_starts = []

    for row_index, row in enumerate(table.rows):
        if not row.cells:
            continue

        cell_text = clean_cell_text(row.cells[0])
        normalized_text = cell_text.strip().upper()

        if ROMAN_RE.fullmatch(normalized_text):
            chapter_starts.append((normalized_text, row_index))

    if len(chapter_starts) != EXPECTED_CHAPTER_COUNT:
        raise ValueError(
            f"Numero capitoli inatteso: trovati {len(chapter_starts)}, "
            f"attesi {EXPECTED_CHAPTER_COUNT}."
        )

    chapters = []

    for i, (roman_number, start_index) in enumerate(chapter_starts):
        if i + 1 < len(chapter_starts):
            end_index = chapter_starts[i + 1][1]
        else:
            end_index = len(table.rows)

        chapter_parts = []

        for row in table.rows[start_index:end_index]:
            if not row.cells:
                continue

            text = clean_cell_text(row.cells[0])

            if text:
                chapter_parts.append(text)

        chapter_text = "\n\n".join(chapter_parts).strip()

        if chapter_text:
            chapters.append(
                {
                    "roman": roman_number,
                    "text": chapter_text,
                }
            )

    return chapters


def load_history(history_path: Path) -> set[str]:
    """
    Legge cronologia.txt e restituisce l'insieme dei capitoli già inviati.
    Se il file non esiste, restituisce un insieme vuoto.
    """
    if not history_path.exists():
        return set()

    sent_chapters = set()

    with history_path.open("r", encoding="utf-8") as file:
        for line in file:
            value = line.strip().upper()

            if value:
                sent_chapters.add(value)

    return sent_chapters


def append_to_history(history_path: Path, roman_number: str) -> None:
    """
    Aggiunge a cronologia.txt il numero romano del capitolo appena inviato.
    """
    with history_path.open("a", encoding="utf-8") as file:
        file.write(roman_number.strip().upper() + "\n")


def choose_random_chapter(chapters: list[dict], history_path: Path) -> dict:
    """
    Sceglie un capitolo random tra quelli non ancora inviati.

    Quando tutti i capitoli sono stati inviati,
    svuota cronologia.txt e ricomincia da capo.
    """
    sent_chapters = load_history(history_path)

    available_chapters = [
        chapter
        for chapter in chapters
        if chapter["roman"].upper() not in sent_chapters
    ]

    if not available_chapters:
        print("Tutti i capitoli sono già stati inviati. Resetto la cronologia.")
        history_path.write_text("", encoding="utf-8")
        available_chapters = chapters

    return random.choice(available_chapters)


def get_required_env(name: str) -> str:
    """
    Legge una variabile d'ambiente obbligatoria.
    Se manca, blocca lo script con un errore chiaro.
    """
    value = os.getenv(name)

    if not value:
        raise EnvironmentError(f"Variabile d'ambiente mancante: {name}")

    return value


def send_email(subject: str, body: str) -> None:
    """
    Invia l'email usando Gmail SMTP SSL.
    Richiede queste variabili d'ambiente:

    - SENDER_EMAIL
    - APP_PASSWORD
    - RECEIVER_EMAIL
    """
    sender_email = get_required_env("SENDER_EMAIL")
    app_password = get_required_env("APP_PASSWORD").replace(" ", "")
    receiver_email = get_required_env("RECEIVER_EMAIL")

    message = EmailMessage()
    message["From"] = sender_email
    message["To"] = receiver_email
    message["Subject"] = subject
    message.set_content(body, subtype="plain", charset="utf-8")

    context = ssl.create_default_context()

    with smtplib.SMTP_SSL(SMTP_SERVER, SMTP_PORT, context=context) as server:
        server.login(sender_email, app_password)
        server.send_message(message)


def main() -> None:
    """
    Funzione principale:
    - controlla se deve inviare ora;
    - carica i capitoli dal DOCX;
    - sceglie un capitolo non ancora inviato;
    - manda l'email;
    - aggiorna cronologia.txt.
    """
    if not should_run_now():
        print("Ora locale non prevista per l'invio. Esco senza inviare.")
        return

    chapters = load_chapters(DOCX_PATH)

    selected_chapter = choose_random_chapter(chapters, HISTORY_PATH)

    roman_number = selected_chapter["roman"]
    chapter_text = selected_chapter["text"]

    subject = f"Spigolature dagli Scritti di Bahá’u’lláh — Capitolo {roman_number}"

    body = (
        f"Capitolo {roman_number}\n\n"
        f"{chapter_text}\n\n"
        "---\n"
        "Invio automatico."
    )

    send_email(subject, body)

    append_to_history(HISTORY_PATH, roman_number)

    print(f"Inviato capitolo {roman_number}.")


if __name__ == "__main__":
    main()
